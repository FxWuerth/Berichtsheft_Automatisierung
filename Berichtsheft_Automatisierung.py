from docx import Document
from datetime import date, timedelta, datetime
import smtplib
from email.message import EmailMessage
import os
import random

user_mail = os.environ["MAIL_USER"]
passwort_mail = os.environ["MAIL_PASSWORD"]

user ={
   'max.musterman@gmail.com': "Max Mustermann",
   
}

#Kalenderberechnungen (Montag_Datum, Freitag_Datum, KV)
heute = date.today()
kalenderwoche = heute.isocalendar().week
wochentag = heute.isoweekday()
# Differenz zu Montag und Freitag berechnen
montag = heute - timedelta(days=wochentag - 1)
freitag = montag + timedelta(days=4)
ausbildungsjahr = 0
if heute > date(2024, 9, 2) and heute < date(2025,9,1):
   ausbildungsjahr = 1
elif  heute >= date(2025,9,1) and heute < date(2026, 8, 31):
   ausbildungsjahr = 2
elif heute >= date(2026, 8, 31):
   ausbildungsjahr = 3

#Berichtsheft Nummer berechnen
ber_num = kalenderwoche+17+(heute.isocalendar().year-2025)*52

#Berufsschul kalenderwochen
beruf = [38, 42, 46, 50, 3, 6, 10, 13, 17, 21, 26, 29]

# zufällige zahlen für die einzelstunden erzeugen
stunden: list[float] =  []
def getrandom(start, end):
    return random.randint(start,end)

def stundnenanzahl():
  ges = 0
  for std in stunden:
     ges += std
  return ges 

def hole_stunden():   
 while True:
   stunden.append(float(getrandom(1,4)))
   if stundnenanzahl() > 7:
       stunden.pop()
   elif len(stunden) == 2 and stundnenanzahl() == 7 or len(stunden) == 2 and stundnenanzahl() < 4:
      stunden.pop() 
   elif stundnenanzahl() < 7 and len(stunden) == 3:
      stunden.pop()
   elif len(stunden) == 3 and stundnenanzahl() == 7:
      break
 randnum = getrandom(0,2)
 stunden[randnum] = stunden[randnum] + 0.2
 stundenstr = [f"{std:.1f}" for std in stunden]
 return stundenstr

#gibt ein Dictionary zurück das beinhaltet ob wir berufsschule haben oder nicht
def Berufsschule(wochen):
    global kalenderwoche 
    if kalenderwoche in wochen:
        return {"[Betrieb]": "Berufsschule"}
    else:
        return{"[Betrieb]": "HRL 3.2"}
    

#Pfade fürs holen und speichern der Word dateien
path = '/app/wordfiles/Berichtsheft_Vorlage.docx'
output = f"/app/wordfiles/Wochenbericht_{ber_num}.docx"


#befüllt die vorlage mit gegeben daten
def fill_berichtsheft(brichtsheft_path, output_path, data):
    doc = Document(brichtsheft_path)
    for table in doc.tables:
     for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    for key, value in data.items():
                      if key in para.text:
                        for run in para.runs:
                          run.text = run.text.replace(key,value)

    doc.save(output_path)

# setzt daten(wichtig für den namen)
def setData(name):
   berufsschule = Berufsschule(beruf)
   temporary = {}
   if berufsschule != {"[Betrieb]": "Berufsschule"}:
    for num in range(5):
     day_hours = hole_stunden()
     stunden.clear()
     for i, h in enumerate(day_hours):
         temporary[f"[ts{i + num*3+1}]"] = h
    tagesstunden = 7.2
    gesamtstunden = 36
   else:
      for num in range(5):
        schule_stunden= ["", "", ""]
        for i, h in enumerate(schule_stunden):
         temporary[f"[ts{i + num*3+1}]"] = h
      tagesstunden = 7.0
      gesamtstunden = 35   

   data = {
         '[Num]': f" {ber_num}",
         '[KV]':  f"KW {kalenderwoche}",
         '[Montag]':  montag.strftime("%d.%m.%Y"),
         '[Freitag]': freitag.strftime("%d.%m.%Y"),
         '[Jahr]': f"{ausbildungsjahr}",
         '[Name]': name,
         '[tgs]': f"{tagesstunden}",
         '[ges]': f"{gesamtstunden}"
    }
   merge = data | berufsschule |temporary
   return merge

#verschieckt eine mail an den reciever
def send_mail(receiver):
# Deine Einstellungen
 smtp_server = "smtp.gmail.com"
 smtp_port = 587
 absender = user_mail
 empfänger = receiver
 passwort = passwort_mail 
 # Betreff und Inhalt
 betreff = f"Berichtsheft KW {kalenderwoche}"
 inhalt = "Im Anhang findest du das Layout für das aktuelle Berichtsheft."

 # Pfad zur Datei
 anhang_pfad = rf"/app/wordfiles/Wochenbericht_{ber_num}.docx"
 
 # E-Mail vorbereiten
 msg = EmailMessage()
 msg["From"] = absender
 msg["To"] = empfänger
 msg["Subject"] = betreff
 msg.set_content(inhalt)

 # Datei anhängen
 with open(anhang_pfad, "rb") as f:
    datei_daten = f.read()
    dateiname = os.path.basename(anhang_pfad)
    msg.add_attachment(datei_daten, maintype="application", subtype="vnd.openxmlformats-officedocument.wordprocessingml.document", filename=dateiname)

 # E-Mail senden
 try:
  with smtplib.SMTP(smtp_server, smtp_port) as server:
    server.starttls()
    server.login(absender, passwort)
    server.send_message(msg)
    print(f"E-Mail erfolgreich gesendet An: {receiver}")
 except Exception as e:
      print(f"Fehaler beim senden: {e}")

# verschickt Mails an alle eingetragenen User
def mail_verteiler():
   for mail, name in user.items():
      fill_berichtsheft(path,output, setData(name))
      send_mail(mail)

#nur wenn das Script direkt ausgeführt wird
if __name__ == '__main__':
    mail_verteiler()
    print(datetime.now())
